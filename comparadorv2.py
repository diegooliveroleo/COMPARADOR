import argparse
import re
import sys
from pathlib import Path
from typing import List, Tuple

import numpy as np
import pandas as pd

# Extracción DOCX
from docx import Document
from unidecode import unidecode

# Stopwords español
import nltk
from nltk.corpus import stopwords

# Métodos clásicos
from rapidfuzz.distance import Levenshtein
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity

# LDA (probabilístico)


# Semántico (SBERT)
from sentence_transformers import SentenceTransformer, util
from sklearn.decomposition import LatentDirichletAllocation
from sklearn.feature_extraction.text import CountVectorizer

# ----------------------------
# Utilidades de carga/limpieza
# ----------------------------
def ensure_nltk_resources():
    try:
        stopwords.words("spanish")
    except LookupError:
        nltk.download("stopwords")


def extract_text_from_docx(path: Path) -> str:
    """Extrae texto de párrafos y tablas de un .docx."""
    doc = Document(path)
    parts: List[str] = []

    for p in doc.paragraphs:
        t = p.text.strip()
        if t:
            parts.append(t)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                t = cell.text.strip()
                if t:
                    parts.append(t)

    return "\n".join(parts)


def normalize_text(text: str) -> str:
    """Normaliza: minúsculas, quita acentos, colapsa espacios."""
    text = text.lower()
    text = unidecode(text)
    text = re.sub(r"\s+", " ", text)
    return text.strip()


def tokenize(text: str, sw: set) -> List[str]:
    """Tokeniza a palabras alfanuméricas simples, filtrando stopwords."""
    tokens = re.findall(r"\b\w+\b", text.lower())
    return [t for t in tokens if t not in sw and len(t) > 1]


def load_docs_from_folder(folder: Path, recursive: bool = False) -> Tuple[List[str], List[str]]:
    pattern = "**/*.docx" if recursive else "*.docx"
    files = sorted(folder.glob(pattern))
    names, texts = [], []
    for f in files:
        try:
            raw = extract_text_from_docx(f)
            names.append(f.name)
            texts.append(raw)
        except Exception as e:
            print(f"[WARN] No se pudo leer {f}: {e}", file=sys.stderr)
    return names, texts


# ----------------------------
# 1) Literal: Levenshtein
# ----------------------------
def similarity_levenshtein(texts: List[str]) -> np.ndarray:
    """Similitud 1 - (distancia / max_len)."""
    n = len(texts)
    norm = [normalize_text(t) for t in texts]
    sims = np.zeros((n, n), dtype=float)
    for i in range(n):
        sims[i, i] = 1.0
        for j in range(i + 1, n):
            a, b = norm[i], norm[j]
            if not a and not b:
                sim = 1.0
            else:
                dist = Levenshtein.distance(a, b)
                denom = max(len(a), len(b)) or 1
                sim = 1.0 - (dist / denom)
            sims[i, j] = sims[j, i] = max(0.0, min(1.0, sim))
    return sims


# ----------------------------
# 2) Conjuntos: Jaccard
# ----------------------------
def similarity_jaccard(texts: List[str], sw: set) -> np.ndarray:
    n = len(texts)
    sets = [set(tokenize(t, sw)) for t in texts]
    sims = np.zeros((n, n), dtype=float)
    for i in range(n):
        sims[i, i] = 1.0
        for j in range(i + 1, n):
            A, B = sets[i], sets[j]
            if not A and not B:
                sim = 1.0
            elif not A or not B:
                sim = 0.0
            else:
                sim = len(A & B) / len(A | B)
            sims[i, j] = sims[j, i] = sim
    return sims


# ----------------------------
# 3) Vectorial clásica: TF-IDF + coseno
# ----------------------------
def similarity_tfidf(texts: List[str], sw_list: List[str]) -> np.ndarray:
    norm_texts = [normalize_text(t) for t in texts]
    vec = TfidfVectorizer(stop_words=sw_list, ngram_range=(1, 2), min_df=1)
    X = vec.fit_transform(norm_texts)
    return cosine_similarity(X)


# ----------------------------
# 4) Probabilístico: LDA + coseno
# ----------------------------
def lda_topic_distributions_sklearn(texts, sw_list, num_topics=None, max_features=5000):
    """
    LDA con scikit-learn:
    - Vectoriza con CountVectorizer (el LDA de sklearn espera conteos).
    - Devuelve matriz d x K con la distribución de temas por documento.
    """
    # Normaliza suavemente (sin quitar tildes aquí no pasa nada)
    cleaned = [re.sub(r"\s+", " ", t.lower()).strip() for t in texts]
    # Vector de conteos (quitando stopwords en español)
    vect = CountVectorizer(stop_words=sw_list, token_pattern=r"\b\w+\b",
                           min_df=2, max_df=0.8, max_features=max_features)
    X = vect.fit_transform(cleaned)

    # Si el vocabulario queda vacío (pocos docs), hacemos un fallback
    if X.shape[1] == 0:
        K = max(2, min(5, len(texts)))
        return np.full((len(texts), K), 1.0 / K)

    # Elegir K (temas)
    if num_topics is None:
        # heurística simple: entre 5 y 20 según tamaño de vocabulario
        K = max(5, min(20, X.shape[1] // 200 or 5))
    else:
        K = max(2, num_topics)

    lda = LatentDirichletAllocation(
        n_components=K,
        learning_method="batch",
        random_state=42,
        max_iter=20,
        n_jobs=-1,
    )
    doc_topic = lda.fit_transform(X)           # d x K, filas normalizadas a 1
    return doc_topic


def similarity_lda(texts: list, sw_set: set, num_topics: int = None) -> np.ndarray:
    """
    Calcula similitud por coseno sobre distribuciones de temas (LDA sklearn).
    """
    sw_list = list(sw_set)
    doc_topic = lda_topic_distributions_sklearn(texts, sw_list, num_topics=num_topics)
    return cosine_similarity(doc_topic)



# ----------------------------
# 5) Semántico: SBERT + coseno
# ----------------------------
def similarity_sbert(texts: List[str], model_name: str = "paraphrase-MiniLM-L6-v2", batch_size: int = 32) -> np.ndarray:
    cleaned = [re.sub(r"\s+", " ", t).strip() for t in texts]
    model = SentenceTransformer(model_name)
    emb = model.encode(
        cleaned,
        batch_size=batch_size,
        convert_to_tensor=True,
        normalize_embeddings=True,
        show_progress_bar=True,
    )
    sim = util.cos_sim(emb, emb).cpu().numpy()
    return sim


# ----------------------------
# Helpers de salida
# ----------------------------
def matrix_to_df(sim: np.ndarray, labels: List[str], scale_percent: bool = True) -> pd.DataFrame:
    df = pd.DataFrame(sim, index=labels, columns=labels)
    if scale_percent:
        df = df * 100.0
    return df.round(2)


def top_pairs(sim_df: pd.DataFrame, top_k: int = 50) -> pd.DataFrame:
    labels = list(sim_df.index)
    n = len(labels)
    rows = []
    for i in range(n):
        for j in range(i + 1, n):
            rows.append({
                "doc_a": labels[i],
                "doc_b": labels[j],
                "sim_%": float(sim_df.iloc[i, j]),
            })
    out = pd.DataFrame(rows).sort_values("sim_%", ascending=False).head(top_k).reset_index(drop=True)
    return out


def main():
    parser = argparse.ArgumentParser(
        description="Compara múltiples .docx aplicando un método por tipo y exporta un Excel."
    )
    parser.add_argument("carpeta", type=str, help="Carpeta con documentos .docx")
    parser.add_argument("--recursivo", action="store_true", help="Incluir subcarpetas")
    parser.add_argument("--modelo", type=str, default="paraphrase-MiniLM-L6-v2", help="Modelo SBERT")
    parser.add_argument("--lda_topics", type=int, default=None, help="Número de temas para LDA (opcional)")
    parser.add_argument("--salida", type=str, default="resultados_similitud.xlsx", help="Nombre del Excel de salida")
    parser.add_argument("--topk", type=int, default=50, help="Nº de pares top por método")
    args = parser.parse_args()

    folder = Path(args.carpeta)
    assert folder.exists() and folder.is_dir(), "La carpeta indicada no existe o no es un directorio."

    ensure_nltk_resources()
    sw_list = stopwords.words("spanish")
    sw_set = set(sw_list)

    print(f"[INFO] Cargando .docx desde: {folder} (recursivo={args.recursivo})")
    names, texts = load_docs_from_folder(folder, recursive=args.recursivo)
    if len(names) < 2:
        print("[ERROR] Se requieren al menos 2 documentos .docx.")
        sys.exit(1)

    print(f"[INFO] Documentos cargados: {len(names)}")
    for n in names:
        print(f" - {n}")

    # Calcular similitudes
    print("\n[INFO] Calculando similitud (1) Literal: Levenshtein...")
    sim_lev = similarity_levenshtein(texts)

    print("[INFO] Calculando similitud (2) Conjuntos: Jaccard...")
    sim_jac = similarity_jaccard(texts, sw_set)

    print("[INFO] Calculando similitud (3) Vectorial: TF-IDF + coseno...")
    sim_tfidf = similarity_tfidf(texts, sw_list)

    print("[INFO] Calculando similitud (4) Probabilístico: LDA + coseno...")
    sim_lda = similarity_lda(texts, sw_set, num_topics=args.lda_topics)

    print("[INFO] Calculando similitud (5) Semántico: SBERT + coseno...")
    sim_sbert = similarity_sbert(texts, model_name=args.modelo)

    # DataFrames (en %)
    df_lev = matrix_to_df(sim_lev, names)
    df_jac = matrix_to_df(sim_jac, names)
    df_tfidf = matrix_to_df(sim_tfidf, names)
    df_lda = matrix_to_df(sim_lda, names)
    df_sbert = matrix_to_df(sim_sbert, names)

    # Top pares por método
    top_lev = top_pairs(df_lev, top_k=args.topk)
    top_jac = top_pairs(df_jac, top_k=args.topk)
    top_tfidf = top_pairs(df_tfidf, top_k=args.topk)
    top_lda = top_pairs(df_lda, top_k=args.topk)
    top_sbert = top_pairs(df_sbert, top_k=args.topk)

    # Escribir Excel
    out_path = Path(args.salida)
    print(f"\n[INFO] Escribiendo Excel: {out_path}")
    with pd.ExcelWriter(out_path, engine="openpyxl") as writer:
        # Matrices
        df_lev.to_excel(writer, sheet_name="1_Literal_Levenshtein")
        df_jac.to_excel(writer, sheet_name="2_Conjuntos_Jaccard")
        df_tfidf.to_excel(writer, sheet_name="3_Vectorial_TFIDF")
        df_lda.to_excel(writer, sheet_name="4_Probabilistico_LDA")
        df_sbert.to_excel(writer, sheet_name="5_Semantico_SBERT")

        # Rankings
        top_lev.to_excel(writer, sheet_name="TOP_Literal", index=False)
        top_jac.to_excel(writer, sheet_name="TOP_Conjuntos", index=False)
        top_tfidf.to_excel(writer, sheet_name="TOP_Vectorial", index=False)
        top_lda.to_excel(writer, sheet_name="TOP_Probabilistico", index=False)
        top_sbert.to_excel(writer, sheet_name="TOP_Semantico", index=False)

        # Info
        info = pd.DataFrame(
            {
                "Parametro": [
                    "Carpeta",
                    "Recursivo",
                    "Modelo SBERT",
                    "LDA topics",
                    "Documentos",
                ],
                "Valor": [
                    str(folder.resolve()),
                    str(args.recursivo),
                    args.modelo,
                    str(args.lda_topics or "auto"),
                    ", ".join(names),
                ],
            }
        )
        info.to_excel(writer, sheet_name="INFO", index=False)

    print("[OK] Finalizado. Hojas creadas:")
    print(" - 1_Literal_Levenshtein")
    print(" - 2_Conjuntos_Jaccard")
    print(" - 3_Vectorial_TFIDF")
    print(" - 4_Probabilistico_LDA")
    print(" - 5_Semantico_SBERT")
    print(" - TOP_* por método y hoja INFO")
    print(f"\nArchivo: {out_path.resolve()}")
    

if __name__ == "__main__":
    main()

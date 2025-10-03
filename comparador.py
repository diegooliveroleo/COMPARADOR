import argparse
import re
from pathlib import Path
from typing import List, Tuple

import numpy as np
import pandas as pd
from docx import Document
from unidecode import unidecode
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity
from sentence_transformers import SentenceTransformer, util


def extract_text_from_docx(path: Path) -> str:
    """Extrae texto de párrafos y tablas de un .docx."""
    doc = Document(path)
    parts: List[str] = []

    # Párrafos
    for p in doc.paragraphs:
        t = p.text.strip()
        if t:
            parts.append(t)

    # Tablas
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                t = cell.text.strip()
                if t:
                    parts.append(t)

    # Nota: no incluye encabezados/pies, comentarios ni objetos incrustados.
    return "\n".join(parts)


def normalize(text: str) -> str:
    """Normaliza a minúsculas, quita acentos y colapsa espacios."""
    text = text.lower()
    text = unidecode(text)
    text = re.sub(r"\s+", " ", text)
    return text.strip()


def load_docs_from_folder(folder: Path, recursive: bool = False) -> Tuple[List[str], List[str]]:
    """Carga todos los .docx y devuelve (nombres, textos)."""
    pattern = "**/*.docx" if recursive else "*.docx"
    files = sorted(folder.glob(pattern))
    names, texts = [], []
    for f in files:
        try:
            raw = extract_text_from_docx(f)
            names.append(f.name)
            texts.append(raw)
        except Exception as e:
            print(f"[WARN] No se pudo leer {f}: {e}")
    return names, texts


def build_tfidf_similarity(texts: List[str]) -> np.ndarray:
    """Similitud coseno con TF-IDF (stopwords español, 1-2gram)."""
    norm_texts = [normalize(t) for t in texts]
    vec = TfidfVectorizer(ngram_range=(1, 2), min_df=1)
    X = vec.fit_transform(norm_texts)
    sim = cosine_similarity(X)
    return sim


def build_semantic_similarity(
    texts: List[str],
    model_name: str = "paraphrase-MiniLM-L6-v2",
    batch_size: int = 32,
) -> np.ndarray:
    """Similitud semántica con SBERT (coseno)."""
    # Para embeddings es útil no eliminar acentos/ñ en exceso; normalizamos suave:
    cleaned = [re.sub(r"\s+", " ", t).strip() for t in texts]
    model = SentenceTransformer(model_name)
    emb = model.encode(
        cleaned,
        batch_size=batch_size,
        convert_to_tensor=True,
        normalize_embeddings=True,  # coseno = producto punto
        show_progress_bar=True,
    )
    sim = util.cos_sim(emb, emb).cpu().numpy()
    return sim


def matrix_to_df(sim: np.ndarray, labels: List[str]) -> pd.DataFrame:
    df = pd.DataFrame(sim, index=labels, columns=labels)
    # Castear a porcentaje con 4 decimales al guardar
    return df


def top_pairs(tfidf_df: pd.DataFrame, sem_df: pd.DataFrame, top_k: int = 50) -> pd.DataFrame:
    """Pares únicos (i<j) ordenados por similitud semántica desc."""
    rows = []
    labels = list(tfidf_df.index)
    n = len(labels)
    for i in range(n):
        for j in range(i + 1, n):
            rows.append({
                "doc_a": labels[i],
                "doc_b": labels[j],
                "sim_tfidf": float(tfidf_df.iloc[i, j]),
                "sim_semantica": float(sem_df.iloc[i, j]),
            })
    df = pd.DataFrame(rows)
    df = df.sort_values("sim_semantica", ascending=False).head(top_k).reset_index(drop=True)
    return df


def save_csv_percent(df: pd.DataFrame, path: Path):
    """Guarda el DataFrame convirtiendo valores numéricos [0,1] a porcentaje."""
    out = df.copy()
    for col in out.columns:
        # Solo intenta convertir si son números
        if np.issubdtype(out[col].dtype, np.number):
            out[col] = (out[col] * 100).round(4)
    out.to_csv(path, encoding="utf-8-sig")


def main():
    parser = argparse.ArgumentParser(
        description="Compara múltiples .docx con TF-IDF y similitud semántica (SBERT)."
    )
    parser.add_argument("carpeta", type=str, help="Carpeta con documentos .docx")
    parser.add_argument("--recursivo", action="store_true", help="Buscar .docx en subcarpetas")
    parser.add_argument("--modelo", type=str, default="paraphrase-MiniLM-L6-v2",
                        help="Modelo SBERT de sentence-transformers")
    parser.add_argument("--topk", type=int, default=10, help="Número de pares más similares a listar")
    args = parser.parse_args()

    folder = Path(args.carpeta)
    assert folder.exists() and folder.is_dir(), "La carpeta indicada no existe o no es un directorio."

    print(f"[INFO] Cargando .docx desde: {folder} (recursivo={args.recursivo})")
    names, texts = load_docs_from_folder(folder, recursive=args.recursivo)

    if len(names) < 2:
        print("[ERROR] Se requieren al menos 2 documentos .docx.")
        return

    print(f"[INFO] Documentos cargados: {len(names)}")
    for n in names:
        print(f" - {n}")

    print("\n[INFO] Calculando similitud TF-IDF...")
    sim_tfidf = build_tfidf_similarity(texts)
    tfidf_df = matrix_to_df(sim_tfidf, names) * 100

    print("\n[INFO] Calculando similitud semántica (SBERT)...")
    sim_sem = build_semantic_similarity(texts, model_name=args.modelo)
    sem_df = matrix_to_df(sim_sem, names) * 100

    # Configuración de impresión legible
    pd.set_option("display.max_columns", None)
    pd.set_option("display.max_rows", None)
    pd.set_option("display.width", 200)

    print("\n========== MATRIZ TF-IDF (%) ==========")
    print(tfidf_df.round(2))

    print("\n========== MATRIZ SEMÁNTICA (%) ==========")
    print(sem_df.round(2))

    # Top pares más similares
    top_df = top_pairs(tfidf_df/100, sem_df/100, top_k=args.topk)
    top_df["sim_tfidf"] = (top_df["sim_tfidf"]*100).round(2)
    top_df["sim_semantica"] = (top_df["sim_semantica"]*100).round(2)

    print(f"\n========== TOP {args.topk} PARES MÁS SIMILARES ==========")
    print(top_df.to_string(index=False))

    print("\n[OK] Comparación finalizada.")



if __name__ == "__main__":
    main()
